"""File handling utilities for Astraeus."""
import os
import logging
import magic
import chardet
from typing import Optional, Tuple, Dict, Any, List
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
import csv
import json
import xml.etree.ElementTree as ET
import yaml
from docx import Document
import openpyxl
from PIL import Image
import mimetypes
import hashlib
from datetime import datetime

from backend.config import ALLOWED_EXTENSIONS, UPLOAD_FOLDER, MAX_FILE_SIZE

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FileProcessingError(Exception):
    """Custom exception for file processing errors."""
    pass

def allowed_file(filename: str) -> bool:
    """
    Check if a file has an allowed extension.
    
    Args:
        filename: Name of the file to check
        
    Returns:
        bool: True if file extension is allowed, False otherwise
    """
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def validate_file(file) -> Tuple[bool, str]:
    """
    Validate a file before processing.
    
    Args:
        file: File object from request.files
        
    Returns:
        Tuple[bool, str]: (is_valid, error_message)
    """
    try:
        # Check if file exists
        if not file:
            return False, "No file provided"
            
        # Check filename
        if not file.filename:
            return False, "No filename provided"
            
        # Check file size
        file.seek(0, os.SEEK_END)
        size = file.tell()
        file.seek(0)
        
        if size > MAX_FILE_SIZE:
            return False, f"File too large. Maximum size is {MAX_FILE_SIZE/1024/1024}MB"
            
        # Check file extension
        if not allowed_file(file.filename):
            return False, f"File type not allowed. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}"
            
        # Check mime type
        mime = magic.Magic(mime=True)
        mime_type = mime.from_buffer(file.read(1024))
        file.seek(0)
        
        if not any(mime_type.startswith(t) for t in ['text/', 'application/pdf', 'application/msword', 
                                                    'application/vnd.openxmlformats-officedocument', 
                                                    'application/vnd.ms-excel']):
            return False, f"Invalid file type: {mime_type}"
            
        return True, ""
        
    except Exception as e:
        logger.error(f"Error validating file: {str(e)}")
        return False, f"Error validating file: {str(e)}"

def save_file(file) -> Tuple[str, str]:
    """
    Save an uploaded file securely with additional validation.
    
    Args:
        file: File object from request.files
        
    Returns:
        Tuple[str, str]: Tuple containing (filename, filepath)
        
    Raises:
        FileProcessingError: If file saving fails
    """
    try:
        # Validate file
        is_valid, error_message = validate_file(file)
        if not is_valid:
            raise FileProcessingError(error_message)
            
        # Generate secure filename
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{timestamp}_{filename}"
        
        # Ensure upload directory exists
        os.makedirs(UPLOAD_FOLDER, exist_ok=True)
        
        # Save file
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(file_path)
        
        # Verify file was saved correctly
        if not os.path.exists(file_path):
            raise FileProcessingError("Failed to save file")
            
        logger.info(f"Saved file: {filename} to {file_path}")
        return filename, file_path
        
    except Exception as e:
        logger.error(f"Error saving file: {str(e)}")
        raise FileProcessingError(f"Error saving file: {str(e)}")

def detect_encoding(file_path: str) -> str:
    """
    Detect file encoding using chardet.
    
    Args:
        file_path: Path to the file
        
    Returns:
        str: Detected encoding
    """
    try:
        with open(file_path, 'rb') as f:
            raw_data = f.read(4096)
            result = chardet.detect(raw_data)
            return result['encoding'] or 'utf-8'
    except Exception as e:
        logger.warning(f"Error detecting encoding for {file_path}: {str(e)}")
        return 'utf-8'

def load_document(file_path: str) -> Optional[str]:
    """
    Load and read the contents of a document with improved error handling.
    Supports multiple file formats including PDF, TXT, CSV, JSON, DOCX, XLSX etc.
    
    Args:
        file_path: Path to the file to read
        
    Returns:
        Optional[str]: Contents of the file if successful, None otherwise
        
    Raises:
        FileProcessingError: If file processing fails
    """
    try:
        if not os.path.exists(file_path):
            raise FileProcessingError(f"File not found: {file_path}")
            
        file_extension = os.path.splitext(file_path)[1].lower()
        encoding = detect_encoding(file_path)
        
        # Text files
        if file_extension in ['.txt', '.md', '.text']:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        
        # PDF files
        elif file_extension == '.pdf':
            try:
                with open(file_path, 'rb') as f:
                    pdf = PdfReader(f)
                    text = []
                    
                    # Extract text from each page
                    for page_num, page in enumerate(pdf.pages, 1):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                # Clean up the extracted text
                                page_text = page_text.replace('\x00', '')  # Remove null bytes
                                page_text = ' '.join(page_text.split())  # Normalize whitespace
                                text.append(f"Page {page_num}:\n{page_text}")
                        except Exception as e:
                            logger.warning(f"Error extracting text from PDF page {page_num}: {str(e)}")
                            continue
                    
                    # If no text was extracted, try alternative method
                    if not text:
                        logger.warning("No text extracted from PDF using standard method, trying alternative...")
                        try:
                            import pytesseract
                            from pdf2image import convert_from_path
                            
                            # Convert PDF to images
                            images = convert_from_path(file_path)
                            for i, image in enumerate(images, 1):
                                # Extract text from image using OCR
                                image_text = pytesseract.image_to_string(image)
                                if image_text:
                                    # Clean up OCR text
                                    image_text = image_text.replace('\x00', '')
                                    image_text = ' '.join(image_text.split())
                                    text.append(f"Page {i} (OCR):\n{image_text}")
                        except Exception as e:
                            logger.error(f"Error in alternative PDF text extraction: {str(e)}")
                    
                    if not text:
                        raise FileProcessingError("Could not extract text from PDF")
                        
                    return '\n\n'.join(text)
                    
            except Exception as e:
                logger.error(f"Error processing PDF file: {str(e)}")
                raise FileProcessingError(f"Error processing PDF: {str(e)}")
        
        # CSV files
        elif file_extension in ['csv']:
            text = []
            with open(file_path, 'r', encoding=encoding) as f:
                csv_reader = csv.reader(f)
                headers = next(csv_reader, None)
                if headers:
                    text.append(', '.join(headers))
                for row in csv_reader:
                    text.append(', '.join(str(cell) for cell in row))
            return '\n'.join(text)
        
        # JSON files
        elif file_extension == 'json':
            with open(file_path, 'r', encoding=encoding) as f:
                data = json.load(f)
                return json.dumps(data, indent=2, ensure_ascii=False)
        
        # XML files
        elif file_extension == 'xml':
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            def process_element(element, indent=0):
                result = [' ' * indent + element.tag]
                if element.text and element.text.strip():
                    result.append(': ' + element.text.strip())
                if element.attrib:
                    result.append(' ' + ' '.join(f'{k}="{v}"' for k, v in element.attrib.items()))
                result.append('\n')
                
                for child in element:
                    result.append(process_element(child, indent + 2))
                return ''.join(result)
            
            return process_element(root)
        
        # YAML files
        elif file_extension in ['yaml', 'yml']:
            with open(file_path, 'r', encoding=encoding) as f:
                data = yaml.safe_load(f)
                return yaml.dump(data, sort_keys=False, default_flow_style=False, allow_unicode=True)

        # Word documents
        elif file_extension == 'docx':
            doc = Document(file_path)
            text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text.append(' | '.join(row_text))
            
            return '\n'.join(text)

        # Excel files
        elif file_extension in ['xlsx', 'xls']:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text = []
            
            for sheet in workbook.sheetnames:
                worksheet = workbook[sheet]
                text.append(f"\nSheet: {sheet}")
                
                # Get column widths for better formatting
                col_widths = []
                for col in worksheet.columns:
                    max_length = 0
                    for cell in col:
                        if cell.value:
                            max_length = max(max_length, len(str(cell.value)))
                    col_widths.append(max_length)
                
                # Process each row
                for row in worksheet.iter_rows():
                    row_values = []
                    for i, cell in enumerate(row):
                        if cell.value is not None:
                            value = str(cell.value)
                            # Pad with spaces for better alignment
                            if i < len(col_widths):
                                value = value.ljust(col_widths[i])
                            row_values.append(value)
                    if row_values:
                        text.append(' | '.join(row_values))
            
            return '\n'.join(text)
        
        # Image files
        elif file_extension in ['jpg', 'jpeg', 'png', 'gif', 'bmp']:
            with Image.open(file_path) as img:
                return f"Image file: {img.format} {img.size} {img.mode}"
        
        else:
            raise FileProcessingError(f"Unsupported file type: {file_extension}")
            
    except Exception as e:
        logger.error(f"Error reading file {file_path}: {str(e)}")
        raise FileProcessingError(f"Error processing file: {str(e)}")

def get_file_info(file_path: str) -> Dict[str, Any]:
    """
    Get detailed information about a file.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Dict[str, Any]: Dictionary containing file metadata
        
    Raises:
        FileProcessingError: If file info retrieval fails
    """
    try:
        if not os.path.exists(file_path):
            raise FileProcessingError(f"File not found: {file_path}")
            
        # Get basic file info
        file_stat = os.stat(file_path)
        file_extension = os.path.splitext(file_path)[1].lower()
        
        # Calculate file hash
        file_hash = hashlib.sha256()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b''):
                file_hash.update(chunk)
        
        # Get mime type
        mime = magic.Magic(mime=True)
        with open(file_path, 'rb') as f:
            mime_type = mime.from_buffer(f.read(1024))
        
        # Basic file info
        info = {
            'filename': os.path.basename(file_path),
            'file_path': file_path,
            'file_size': file_stat.st_size,
            'file_type': mime_type,
            'extension': file_extension,
            'created_at': datetime.fromtimestamp(file_stat.st_ctime).isoformat(),
            'modified_at': datetime.fromtimestamp(file_stat.st_mtime).isoformat(),
            'sha256_hash': file_hash.hexdigest()
        }
        
        # Add format-specific metadata
        if file_extension == '.pdf':
            try:
                with open(file_path, 'rb') as f:
                    pdf = PdfReader(f)
                    info.update({
                        'page_count': len(pdf.pages),
                        'is_encrypted': pdf.is_encrypted,
                        'metadata': pdf.metadata
                    })
            except Exception as e:
                logger.warning(f"Error getting PDF metadata: {str(e)}")
                
        elif file_extension in ['.docx', '.doc']:
            try:
                doc = Document(file_path)
                info.update({
                    'paragraph_count': len(doc.paragraphs),
                    'table_count': len(doc.tables)
                })
            except Exception as e:
                logger.warning(f"Error getting DOCX metadata: {str(e)}")
                
        elif file_extension in ['.xlsx', '.xls']:
            try:
                wb = openpyxl.load_workbook(file_path, read_only=True)
                info.update({
                    'sheet_count': len(wb.sheetnames),
                    'sheet_names': wb.sheetnames
                })
            except Exception as e:
                logger.warning(f"Error getting Excel metadata: {str(e)}")
                
        elif file_extension in ['.jpg', '.jpeg', '.png', '.gif']:
            try:
                with Image.open(file_path) as img:
                    info.update({
                        'image_size': img.size,
                        'image_mode': img.mode,
                        'image_format': img.format
                    })
            except Exception as e:
                logger.warning(f"Error getting image metadata: {str(e)}")
        
        return info
        
    except Exception as e:
        logger.error(f"Error getting file info for {file_path}: {str(e)}")
        raise FileProcessingError(f"Error getting file metadata: {str(e)}")

def cleanup_file(file_path: str) -> bool:
    """
    Safely remove a file and its associated resources.
    
    Args:
        file_path: Path to the file to remove
        
    Returns:
        bool: True if cleanup was successful, False otherwise
    """
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
            logger.info(f"Removed file: {file_path}")
            return True
        return False
    except Exception as e:
        logger.error(f"Error cleaning up file {file_path}: {str(e)}")
        return False 